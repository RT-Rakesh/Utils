# -*- coding: utf-8 -*-
"""
Created on Wen Dec 18 2018
@author: Rt-Rakesh

This script is a utility to plot the precision and recall for the object_detection model.
Usage: Should be used by calling in jupyter notbooks.
"""
import numpy as np
import tensorflow as tf
import matplotlib.pyplot as plt
import os
import itertools
from docx import Document

# sys.path.append('/home/tools/tensorflow_object_detection_api')
from object_detection.core import standard_fields
from object_detection.metrics import tf_example_parser
from object_detection.utils import label_map_util


def compute_iou(groundtruth_box, detection_box):
    g_ymin, g_xmin, g_ymax, g_xmax = tuple(groundtruth_box.tolist())
    d_ymin, d_xmin, d_ymax, d_xmax = tuple(detection_box.tolist())

    xa = max(g_xmin, d_xmin)
    ya = max(g_ymin, d_ymin)
    xb = min(g_xmax, d_xmax)
    yb = min(g_ymax, d_ymax)

    intersection = max(0, xb - xa + 1) * max(0, yb - ya + 1)

    boxAArea = (g_xmax - g_xmin + 1) * (g_ymax - g_ymin + 1)
    boxBArea = (d_xmax - d_xmin + 1) * (d_ymax - d_ymin + 1)

    return intersection / float(boxAArea + boxBArea - intersection)


def process_detections(detections_record,
                       categories,
                       iou_threshold,
                       confidence_threshold):

    IOU_THRESHOLD = iou_threshold
    CONFIDENCE_THRESHOLD = confidence_threshold
    record_iterator = tf.python_io.tf_record_iterator(path=detections_record)
    data_parser = tf_example_parser.TfExampleDetectionAndGTParser()

    # confusion_matrix = np.zeros(shape=(len(categories) + 1, len(categories) + 1))
    confusion_matrix = np.zeros(shape=(len(categories) + 2, len(categories) + 2))

    image_index = 0
    for string_record in record_iterator:
        example = tf.train.Example()
        example.ParseFromString(string_record)
        decoded_dict = data_parser.parse(example)

        image_index += 1

        if decoded_dict:
            groundtruth_boxes = decoded_dict[standard_fields.InputDataFields.groundtruth_boxes]
            groundtruth_classes = decoded_dict[standard_fields.InputDataFields.groundtruth_classes]

            detection_scores = decoded_dict[standard_fields.DetectionResultFields.detection_scores]
            detection_classes = decoded_dict[standard_fields.DetectionResultFields.detection_classes][detection_scores >= CONFIDENCE_THRESHOLD]
            detection_boxes = decoded_dict[standard_fields.DetectionResultFields.detection_boxes][detection_scores >= CONFIDENCE_THRESHOLD]

            matches = []

            if image_index % 100 == 0:
                print("Processed %d images" % (image_index))

            for i in range(len(groundtruth_boxes)):
                for j in range(len(detection_boxes)):
                    iou = compute_iou(groundtruth_boxes[i], detection_boxes[j])

                    if iou > IOU_THRESHOLD:
                        matches.append([i, j, iou])

            matches = np.array(matches)
            if matches.shape[0] > 0:
                # Sort list of matches by descending IOU so we can remove duplicate detections
                # while keeping the highest IOU entry.
                matches = matches[matches[:, 2].argsort()[::-1][:len(matches)]]

                # Remove duplicate detections from the list.
                matches = matches[np.unique(matches[:, 1], return_index=True)[1]]

                # Sort the list again by descending IOU. Removing duplicates doesn't preserve
                # our previous sort.
                matches = matches[matches[:, 2].argsort()[::-1][:len(matches)]]

                # Remove duplicate ground truths from the list.
                matches = matches[np.unique(matches[:, 0], return_index=True)[1]]

            for i in range(len(groundtruth_boxes)):
                if matches.shape[0] > 0 and matches[matches[:, 0] == i].shape[0] == 1:
                    confusion_matrix[groundtruth_classes[i] - 1][detection_classes[int(matches[matches[:, 0] == i, 1][0])] - 1] += 1
                else:
                    confusion_matrix[groundtruth_classes[i] - 1][confusion_matrix.shape[1] - 2] += 1

                confusion_matrix[groundtruth_classes[i] - 1][confusion_matrix.shape[1] - 1] += 1

            for i in range(len(detection_boxes)):
                if matches.shape[0] > 0 and matches[matches[:, 1] == i].shape[0] == 0:
                    confusion_matrix[confusion_matrix.shape[0] - 2][detection_classes[i] - 1] += 1

                confusion_matrix[confusion_matrix.shape[0] - 1][detection_classes[i] - 1] += 1
        else:
            print("Skipped image %d" % (image_index))

    print("Processed %d images" % (image_index))

    return confusion_matrix


def display(confusion_matrix,
            categories,
            iou_threshold):
    IOU_THRESHOLD = iou_threshold
    print("\nConfusion Matrix:")
    print(confusion_matrix, "\n")

    for i in range(len(categories)):
        id = categories[i]["id"] - 1
        name = categories[i]["name"]

        total_target = np.sum(confusion_matrix[id, :])
        total_predicted = np.sum(confusion_matrix[:, id])

        precision = float(confusion_matrix[id, id] / total_predicted)
        recall = float(confusion_matrix[id, id] / total_target)

        print('precision_{}@{}IOU: {:.2f}'.format(name, IOU_THRESHOLD, precision))
        print('recall_{}@{}IOU: {:.2f}'.format(name, IOU_THRESHOLD, recall))


def generate_prec_recall_doc(confusion_matrix,
                             categories,
                             iou_threshold,
                             save_path):
    IOU_THRESHOLD = iou_threshold

    label_list = []
    for i in range(len(categories)):
        label_list.append(categories[i]["name"])
    record = []
    for i in range(len(categories)):
        id = categories[i]["id"] - 1
        name = categories[i]["name"]
        total_target = np.sum(confusion_matrix[id, :])
        total_predicted = np.sum(confusion_matrix[:, id])
        precision = float(confusion_matrix[id, id] / total_predicted)
        recall = float(confusion_matrix[id, id] / total_target)
        f1 = 2*(recall * precision) / (recall + precision)
        # print('precision_{}@{}IOU: {:.2f}'.format(name, IOU_THRESHOLD, precision))
        # print('recall_{}@{}IOU: {:.2f}'.format(name, IOU_THRESHOLD, recall))
        record.append((name, round(precision, 2), round(recall, 2), round(f1, 2)))
    record = (record)
    document = Document()
    document.add_heading('Precision & Recall Label Wise ', 0)
    p = document.add_paragraph('This model has handled the following lables ')
    for t in label_list:
        p.add_run(t+', ').bold = True
    p.add_run('.')
    p.add_run('')
    p.add_run('')
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Label').bold = True
    hdr_cells[1].paragraphs[0].add_run('precision@{}IOU'.format(IOU_THRESHOLD)).bold = True
    hdr_cells[2].paragraphs[0].add_run('recall@{}IOU'.format(IOU_THRESHOLD)).bold = True
    hdr_cells[3].paragraphs[0].add_run('f1@{}IOU'.format(IOU_THRESHOLD)).bold = True
    for la, precision, recall, f1 in record:
        row_cells = table.add_row().cells
        row_cells[0].text = str(la)
        row_cells[1].text = str(precision)
        row_cells[2].text = str(recall)
        row_cells[3].text = str(f1)
    document.add_page_break()
    document.save(os.path.join(save_path, 'precision_recall.docx'))
    print("Precsion & Recall document generated Successfully")


def plot_confusion_matrix(cm,
                          target_names,
                          image_path,
                          title='Confusion matrix',
                          cmap=None,
                          normalize=True):
    """
    given a sklearn confusion matrix (cm), make a nice plot

    Arguments
    ---------
    cm:           confusion matrix from sklearn.metrics.confusion_matrix

    target_names: given classification classes such as [0, 1, 2]
                  the class names, for example: ['high', 'medium', 'low']

    title:        the text to display at the top of the matrix

    cmap:         the gradient of the values displayed from matplotlib.pyplot.cm
                  see http://matplotlib.org/examples/color/colormaps_reference.html
                  plt.get_cmap('jet') or plt.cm.Blues

    normalize:    If False, plot the raw numbers
                  If True, plot the proportions

    Usage
    -----
    plot_confusion_matrix(cm           = cm,                  # confusion matrix created by
                                                              # sklearn.metrics.confusion_matrix
                          normalize    = True,                # show proportions
                          target_names = y_labels_vals,       # list of names of the classes
                          title        = best_estimator_name) # title of graph

    Citiation
    ---------
    http://scikit-learn.org/stable/auto_examples/model_selection/plot_confusion_matrix.html

    """

    accuracy = np.trace(cm) / float(np.sum(cm))
    misclass = 1 - accuracy

    if cmap is None:
        cmap = plt.get_cmap('Blues')

    plt.figure(figsize=(15, 11))
    plt.imshow(cm, interpolation='nearest', cmap=cmap)
    plt.title(title)
    plt.colorbar()

    if target_names is not None:
        tick_marks = np.arange(len(target_names))
        plt.xticks(tick_marks, target_names, rotation=45)
        plt.yticks(tick_marks, target_names)

    if normalize:
        cm = cm.astype('float') / cm.sum(axis=1)[:, np.newaxis]

    thresh = cm.max() / 1.5 if normalize else cm.max() / 2
    for i, j in itertools.product(range(cm.shape[0]), range(cm.shape[1])):
        if normalize:
            plt.text(j, i, "{:0.4f}".format(cm[i, j]),
                     horizontalalignment="center",
                     color="white" if cm[i, j] > thresh else "black")
        else:
            plt.text(j, i, "{:,}".format(cm[i, j]),
                     horizontalalignment="center",
                     color="white" if cm[i, j] > thresh else "black")

    plt.tight_layout()
    plt.ylabel('True label')
    plt.xlabel('Predicted label\n\nFor True Label: "Nothing" Objects are part of the ground-truth but weren’t detected are counted\n in the last column of the matrix (in the row corresponding to the ground-truth class).\n\nFor Predicted: "Nothing"Objects were detected but aren’t part of the confusion matrix are counted\n in the last row of the matrix (in the column corresponding to the detected class).'.format(accuracy, misclass))
    plt.tight_layout()
    plt.savefig(os.path.join(image_path, 'Confusion_Matrix.png'), dpi=200)
    plt.close()
    print("Confusion Matrix generated Successfully")


def generate_confusion_matrix(odapi_path,
                              label_map_path,
                              save_path,
                              detection_tf_record_path,
                              iou_threshold=0.5,
                              confidence_threshold=0.5,
                              display_cm=False):
    """
    This Function generates the confusion matrix, presicion recall document for all the labels handles by the moadel.
    -----
    Args:
    -----
    1.odapi_path: --str The object detection api source code path( This is optional)
    2.label_map_path: --str  The path to the label map.
    3.detection_tf_record_path: -- str The path to the detections tf record path.
    4.iou_threshold: --float Deault:0.5
    5.confidence_threshold: --float Default:0.5
    6.display_cm: --Boolean Deafault:False Displays the confusion matrix, the label wise presicion and recall.
    ---------
    Returns:
    ---------
    Saves confusion matrix image and precision recall document.
    """
    # import_object_detection_api(odapi_path)
    label_map = label_map_util.load_labelmap(label_map_path)
    categories = label_map_util.convert_label_map_to_categories(label_map, max_num_classes=100, use_display_name=True)
    confusion_matrix = process_detections(detection_tf_record_path, categories, iou_threshold, confidence_threshold)
    if display_cm:
        display(confusion_matrix, categories, iou_threshold)
    name_list = []
    for i in range(len(categories)):
        name_list.append(categories[i]["name"])
    name_list.append('Nothing')
    name_list.append('Total')
    target_names = name_list
    plot_confusion_matrix(confusion_matrix, target_names, save_path, title='Confusion matrix', cmap=None, normalize=False)
    generate_prec_recall_doc(confusion_matrix, categories, iou_threshold, save_path)
